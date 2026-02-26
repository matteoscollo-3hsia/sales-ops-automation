from __future__ import annotations

import json
import sys
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn


def _length_to_pt(value) -> float | None:
    if value is None:
        return None
    try:
        return value.pt
    except Exception:
        return None


def _style_info(style) -> dict[str, Any]:
    font = style.font
    fmt = style.paragraph_format
    return {
        "font_name": font.name,
        "font_size_pt": _length_to_pt(font.size),
        "bold": font.bold,
        "spacing_before_pt": _length_to_pt(fmt.space_before),
        "spacing_after_pt": _length_to_pt(fmt.space_after),
        "keep_with_next": fmt.keep_with_next,
    }


def _get_table_width_info(table) -> dict[str, Any]:
    tbl_pr = table._tbl.tblPr
    tblw = tbl_pr.find(qn("w:tblW")) if tbl_pr is not None else None
    layout = tbl_pr.find(qn("w:tblLayout")) if tbl_pr is not None else None
    width_type = tblw.get(qn("w:type")) if tblw is not None else None
    width = tblw.get(qn("w:w")) if tblw is not None else None
    layout_type = layout.get(qn("w:type")) if layout is not None else None
    return {"width_type": width_type, "width": width, "layout": layout_type}


def _get_table_column_widths(table) -> list[float | None]:
    widths: list[float | None] = []
    for col in table.columns:
        widths.append(_length_to_pt(col.width))
    if any(width is not None for width in widths):
        return widths
    if table.rows:
        first_row = table.rows[0]
        widths = [_length_to_pt(cell.width) for cell in first_row.cells]
    return widths


def _row_has_header_repeat(row) -> bool:
    tr_pr = row._tr.trPr
    if tr_pr is None:
        return False
    return tr_pr.find(qn("w:tblHeader")) is not None


def _detect_list_paragraphs(doc: Document) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        has_num = bool(paragraph._p.xpath(".//w:numPr"))
        if (
            has_num
            or "List" in style_name
            or "Bullet" in style_name
            or "Number" in style_name
        ):
            fmt = paragraph.paragraph_format
            items.append(
                {
                    "style": style_name,
                    "left_indent_pt": _length_to_pt(fmt.left_indent),
                    "first_line_indent_pt": _length_to_pt(fmt.first_line_indent),
                }
            )
    return items


def _select_body_style(style_counts: Counter[str]) -> list[str]:
    excluded = {"Title", "Heading 1", "Heading 2", "Heading 3"}
    candidates: list[tuple[str, int]] = []
    for name, count in style_counts.items():
        if name in excluded:
            continue
        if "List" in name or "Bullet" in name or "Number" in name:
            continue
        candidates.append((name, count))
    candidates.sort(key=lambda item: item[1], reverse=True)
    return [name for name, _ in candidates[:3]]


def extract_docx_spec(docx_path: Path) -> dict[str, Any]:
    doc = Document(str(docx_path))
    style_counts: Counter[str] = Counter()
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else "None"
        style_counts[style_name] += 1

    paragraph_styles_used = [
        {"style": name, "count": count} for name, count in style_counts.most_common()
    ]

    key_styles = ["Title", "Heading 1", "Heading 2", "Heading 3", "Normal"]
    body_candidates = _select_body_style(style_counts)
    for name in body_candidates:
        if name not in key_styles:
            key_styles.append(name)

    style_details: dict[str, Any] = {}
    for style_name in key_styles:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style_details[style_name] = _style_info(style)

    tables = []
    for table in doc.tables:
        width_info = _get_table_width_info(table)
        header_repeat = False
        for row in table.rows[:1]:
            header_repeat = _row_has_header_repeat(row)
        tables.append(
            {
                "rows": len(table.rows),
                "cols": len(table.columns),
                "width_type": width_info.get("width_type"),
                "width": width_info.get("width"),
                "layout": width_info.get("layout"),
                "column_widths_pt": _get_table_column_widths(table),
                "header_repeat_first_row": header_repeat,
            }
        )

    list_paragraphs = _detect_list_paragraphs(doc)

    return {
        "paragraph_styles_used": paragraph_styles_used,
        "style_details": style_details,
        "body_style_candidates": body_candidates,
        "tables": {
            "count": len(tables),
            "items": tables,
        },
        "lists": {
            "count": len(list_paragraphs),
            "items": list_paragraphs,
        },
    }


def main() -> int:
    if len(sys.argv) < 2:
        print(
            "Usage: python scripts/extract_docx_spec.py path/to/docx", file=sys.stderr
        )
        return 2
    docx_path = Path(sys.argv[1])
    if not docx_path.exists():
        print(f"Missing docx: {docx_path}", file=sys.stderr)
        return 2

    spec = extract_docx_spec(docx_path)
    out_dir = Path("out")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "docx_spec.json"
    out_path.write_text(json.dumps(spec, indent=2), encoding="utf-8")
    print(f"Wrote: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
