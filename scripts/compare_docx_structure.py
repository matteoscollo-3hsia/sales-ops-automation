from __future__ import annotations

import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn


def _row_has_header_repeat(row) -> bool:
    tr_pr = row._tr.trPr
    if tr_pr is None:
        return False
    return tr_pr.find(qn("w:tblHeader")) is not None


def _table_summary(doc: Document) -> list[dict[str, Any]]:
    tables = []
    for table in doc.tables:
        header_repeat = False
        if table.rows:
            header_repeat = _row_has_header_repeat(table.rows[0])
        tables.append(
            {
                "rows": len(table.rows),
                "cols": len(table.columns),
                "header_repeat_first_row": header_repeat,
            }
        )
    return tables


def _heading_counts(doc: Document) -> dict[str, int]:
    counts = {"Title": 0, "Heading 1": 0, "Heading 2": 0, "Heading 3": 0}
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name in counts:
            counts[style_name] += 1
    return counts


def _is_separator(paragraph) -> bool:
    text = paragraph.text.strip()
    if text in ("---", "—", "–", "***"):
        return True
    borders = paragraph._p.xpath(".//w:pBdr")
    if borders:
        return True
    return False


def _separator_count(doc: Document) -> int:
    return sum(1 for paragraph in doc.paragraphs if _is_separator(paragraph))


def _is_list_paragraph(paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    has_num = bool(paragraph._p.xpath(".//w:numPr"))
    return (
        has_num
        or "List" in style_name
        or "Bullet" in style_name
        or "Number" in style_name
    )


def _list_count(doc: Document) -> int:
    return sum(1 for paragraph in doc.paragraphs if _is_list_paragraph(paragraph))


@dataclass
class CompareResult:
    report_path: Path
    major_mismatches: list[str]
    minor_notes: list[str]


def compare_docx_structure(
    generated_path: Path, golden_path: Path, report_path: Path
) -> CompareResult:
    gen_doc = Document(str(generated_path))
    golden_doc = Document(str(golden_path))

    major: list[str] = []
    minor: list[str] = []

    gen_tables = _table_summary(gen_doc)
    golden_tables = _table_summary(golden_doc)

    if len(gen_tables) != len(golden_tables):
        major.append(
            f"Table count mismatch: generated={len(gen_tables)} golden={len(golden_tables)}"
        )

    table_lines = []
    for idx, golden_table in enumerate(golden_tables):
        if idx >= len(gen_tables):
            table_lines.append(f"- Table {idx + 1}: missing in generated output")
            continue
        gen_table = gen_tables[idx]
        row_diff = abs(gen_table["rows"] - golden_table["rows"])
        cols_match = gen_table["cols"] == golden_table["cols"]
        if not cols_match:
            minor.append(
                f"Table {idx + 1} column mismatch: generated={gen_table['cols']} golden={golden_table['cols']}"
            )
        if row_diff > 2:
            minor.append(
                f"Table {idx + 1} row mismatch: generated={gen_table['rows']} golden={golden_table['rows']}"
            )
        if (
            golden_table["header_repeat_first_row"]
            and not gen_table["header_repeat_first_row"]
        ):
            major.append(f"Table {idx + 1} missing header-repeat on first row")
        table_lines.append(
            f"- Table {idx + 1}: rows {gen_table['rows']} vs {golden_table['rows']}, "
            f"cols {gen_table['cols']} vs {golden_table['cols']}, "
            f"header-repeat gen={gen_table['header_repeat_first_row']} "
            f"golden={golden_table['header_repeat_first_row']}"
        )

    heading_gen = _heading_counts(gen_doc)
    heading_gold = _heading_counts(golden_doc)

    separator_gen = _separator_count(gen_doc)
    separator_gold = _separator_count(golden_doc)

    list_gen = _list_count(gen_doc)
    list_gold = _list_count(golden_doc)

    report_lines = [
        "# DOCX Compare Report",
        "",
        f"Generated: `{generated_path}`",
        f"Golden: `{golden_path}`",
        "",
        "## Tables",
    ]
    report_lines.extend(table_lines or ["- No tables detected"])
    report_lines.extend(
        [
            "",
            "## Headings",
            f"- Title: generated={heading_gen['Title']} golden={heading_gold['Title']}",
            f"- Heading 1: generated={heading_gen['Heading 1']} golden={heading_gold['Heading 1']}",
            f"- Heading 2: generated={heading_gen['Heading 2']} golden={heading_gold['Heading 2']}",
            f"- Heading 3: generated={heading_gen['Heading 3']} golden={heading_gold['Heading 3']}",
            "",
            "## Separators",
            f"- Separators: generated={separator_gen} golden={separator_gold}",
            "",
            "## Lists",
            f"- List paragraphs: generated={list_gen} golden={list_gold}",
            "",
            "## Notes",
        ]
    )
    if major:
        report_lines.append("Major mismatches:")
        report_lines.extend([f"- {item}" for item in major])
    if minor:
        report_lines.append("Minor notes:")
        report_lines.extend([f"- {item}" for item in minor])
    if not major and not minor:
        report_lines.append("- No structural differences detected.")

    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text("\n".join(report_lines) + "\n", encoding="utf-8")
    return CompareResult(
        report_path=report_path, major_mismatches=major, minor_notes=minor
    )


def main() -> int:
    if len(sys.argv) < 3:
        print(
            "Usage: python scripts/compare_docx_structure.py generated.docx golden.docx",
            file=sys.stderr,
        )
        return 2
    generated_path = Path(sys.argv[1])
    golden_path = Path(sys.argv[2])
    if not generated_path.exists():
        print(f"Missing generated docx: {generated_path}", file=sys.stderr)
        return 2
    if not golden_path.exists():
        print(f"Missing golden docx: {golden_path}", file=sys.stderr)
        return 2

    report_path = Path("out") / "docx_compare_report.md"
    result = compare_docx_structure(generated_path, golden_path, report_path)
    print(f"Wrote: {result.report_path}")
    if result.major_mismatches:
        return 3
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
