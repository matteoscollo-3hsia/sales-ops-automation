from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document


def summarize_docx_runs(docx_path: str | Path) -> dict[str, Any]:
    path = Path(docx_path)
    doc = Document(str(path))
    summary: dict[str, Any] = {
        "docx_path": str(path),
        "total_runs": 0,
        "bold_runs": 0,
        "italic_runs": 0,
        "code_runs": 0,
        "paragraph_text": [],
        "table_cell_text": [],
    }

    for paragraph in doc.paragraphs:
        summary["paragraph_text"].append(paragraph.text)
        _accumulate_run_counts(summary, paragraph.runs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                summary["table_cell_text"].append(cell.text)
                for paragraph in cell.paragraphs:
                    _accumulate_run_counts(summary, paragraph.runs)

    return summary


def assert_run_thresholds(
    summary: dict[str, Any],
    *,
    min_bold: int = 0,
    min_italic: int = 0,
    min_code: int = 0,
) -> None:
    failures: list[str] = []
    if summary.get("bold_runs", 0) < min_bold:
        failures.append(
            f"bold_runs={summary.get('bold_runs', 0)} is below required minimum {min_bold}"
        )
    if summary.get("italic_runs", 0) < min_italic:
        failures.append(
            f"italic_runs={summary.get('italic_runs', 0)} is below required minimum {min_italic}"
        )
    if summary.get("code_runs", 0) < min_code:
        failures.append(
            f"code_runs={summary.get('code_runs', 0)} is below required minimum {min_code}"
        )
    if failures:
        raise AssertionError("; ".join(failures))


def assert_expected_substrings(summary: dict[str, Any], expected: list[str]) -> None:
    if not expected:
        return
    haystack = "\n".join(summary.get("paragraph_text", []) + summary.get("table_cell_text", []))
    missing = [snippet for snippet in expected if snippet not in haystack]
    if missing:
        raise AssertionError(f"Missing expected substrings: {missing}")


def _accumulate_run_counts(summary: dict[str, Any], runs) -> None:
    for run in runs:
        summary["total_runs"] += 1
        if bool(run.bold):
            summary["bold_runs"] += 1
        if bool(run.italic):
            summary["italic_runs"] += 1
        if (run.font.name or "").strip().lower() == "consolas":
            summary["code_runs"] += 1


def main() -> int:
    parser = argparse.ArgumentParser(description="Dump and assert DOCX run styling counts.")
    parser.add_argument("docx_path", help="Path to .docx file")
    parser.add_argument("--min-bold", type=int, default=0, help="Minimum bold runs required")
    parser.add_argument(
        "--min-italic", type=int, default=0, help="Minimum italic runs required"
    )
    parser.add_argument("--min-code", type=int, default=0, help="Minimum code runs required")
    parser.add_argument(
        "--expect-substring",
        action="append",
        default=[],
        help="Substring expected in rendered paragraph/cell text (repeatable)",
    )
    args = parser.parse_args()

    summary = summarize_docx_runs(args.docx_path)
    print(json.dumps(summary, indent=2, ensure_ascii=True))

    try:
        assert_run_thresholds(
            summary,
            min_bold=args.min_bold,
            min_italic=args.min_italic,
            min_code=args.min_code,
        )
        assert_expected_substrings(summary, args.expect_substring)
    except AssertionError as exc:
        print(f"ASSERTION FAILED: {exc}")
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
