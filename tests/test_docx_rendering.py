from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE

from primer_ops.render_docx import render_primer_docx


def _style_exists(doc: Document, name: str) -> bool:
    try:
        doc.styles[name]
    except KeyError:
        return False
    return True


def _get_paragraph_by_text(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise AssertionError(f"Paragraph not found: {text}")


def test_heading_styles(tmp_path: Path) -> None:
    md_path = tmp_path / "sample.md"
    docx_path = tmp_path / "sample.docx"
    md_path.write_text(
        "# Title\n\n## Section\n\n### Subsection\n\nParagraph.\n", encoding="utf-8"
    )
    render_primer_docx(str(md_path), str(docx_path), None)
    doc = Document(str(docx_path))

    title_para = _get_paragraph_by_text(doc, "Title")
    section_para = _get_paragraph_by_text(doc, "Section")
    subsection_para = _get_paragraph_by_text(doc, "Subsection")

    expected_title = "Title" if _style_exists(doc, "Title") else "Heading 1"
    if _style_exists(doc, expected_title):
        assert title_para.style.name == expected_title, (
            f"Title style mismatch: {title_para.style.name}"
        )

    if _style_exists(doc, "Heading 1"):
        assert section_para.style.name == "Heading 1", (
            f"Heading 1 style mismatch: {section_para.style.name}"
        )

    if _style_exists(doc, "Heading 2"):
        assert subsection_para.style.name == "Heading 2", (
            f"Heading 2 style mismatch: {subsection_para.style.name}"
        )


def test_inline_markdown_render(tmp_path: Path) -> None:
    md_path = tmp_path / "inline.md"
    docx_path = tmp_path / "inline.docx"
    md_path.write_text(
        "Paragraph with **bold** and *italic* and [site](https://example.com) "
        "plus ([site](https://example.com)).\n",
        encoding="utf-8",
    )
    render_primer_docx(str(md_path), str(docx_path), None)
    doc = Document(str(docx_path))
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "**" not in full_text, "Raw bold markers leaked into DOCX text."
    assert "](" not in full_text, "Raw link markers leaked into DOCX text."
    doc_xml = doc.element.xml
    assert "site" in full_text or "site" in doc_xml, "Expected link label text missing."

    has_url_text = "https://example.com" in full_text
    has_hyperlink_rel = any(
        rel.reltype == RELATIONSHIP_TYPE.HYPERLINK
        and rel.target_ref == "https://example.com"
        for rel in doc.part.rels.values()
    )
    assert has_url_text or has_hyperlink_rel, "Expected link URL not found."


def test_numeric_heading_normalization(tmp_path: Path) -> None:
    md_path = tmp_path / "numeric.md"
    docx_path = tmp_path / "numeric.docx"
    md_path.write_text(
        "1. Company Introduction\n\n2.1 Ownership & Governance\n\nParagraph.\n",
        encoding="utf-8",
    )
    render_primer_docx(str(md_path), str(docx_path), None)
    doc = Document(str(docx_path))
    intro_para = _get_paragraph_by_text(doc, "1. Company Introduction")
    ownership_para = _get_paragraph_by_text(doc, "2.1 Ownership & Governance")

    if _style_exists(doc, "Heading 1"):
        assert intro_para.style.name == "Heading 1", (
            f"Numeric heading styled as {intro_para.style.name}"
        )
    else:
        assert intro_para.style.name != "List Number"

    if _style_exists(doc, "Heading 2"):
        assert ownership_para.style.name == "Heading 2", (
            f"Numeric heading styled as {ownership_para.style.name}"
        )
    else:
        assert ownership_para.style.name != "List Number"
