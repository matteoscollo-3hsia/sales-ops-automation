from __future__ import annotations

from pathlib import Path
from datetime import date
import json
import re
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt
from docx.text.paragraph import Paragraph

PLACEHOLDER = "{{CONTENT}}"
_NUMERIC_HEADING_RE = re.compile(r"^(?P<num>\d+(?:\.\d+)*)(?:\.)?\s+\S")
_ORDERED_LIST_RE = re.compile(r"^\d+\.\s+\S")
_TABLE_SEP_LINE_RE = re.compile(r"^[\s\|\-:]+$")
_TABLE_SEP_CELL_RE = re.compile(r"^:?-{3,}:?$")
_FENCE_RE = re.compile(r"^\s*(```|~~~)")
_REQUIRE_NUMERIC_HEADING_CAPS = True


def render_primer_docx(
    md_path: str, out_docx_path: str, template_path: str | None = None
) -> None:
    md_file = Path(md_path)
    out_file = Path(out_docx_path)
    template_file = Path(template_path) if template_path else None

    markdown_text = md_file.read_text(encoding="utf-8")
    markdown_text = normalize_markdown_for_docx(markdown_text)
    doc, writer = _init_document(template_file)
    _apply_placeholder_replacements(doc, md_file)
    _apply_style_profile(doc)
    parser = _markdown_parser()
    parser_name = "markdown_it" if _is_markdown_it(parser) else "fallback"
    print(f"DOCX parser: {parser_name}")
    tokens = parser.parse(markdown_text)
    _render_tokens(tokens, doc, writer)

    out_file.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_atomic(doc, out_file)


def normalize_markdown_for_docx(text: str) -> str:
    if not text:
        return text
    has_trailing_newline = text.endswith("\n")
    lines = text.splitlines()
    fence_mask = _compute_fence_mask(lines)
    lines = _normalize_table_separators(lines, fence_mask)
    table_mask = _compute_table_mask(lines, fence_mask)
    lines = _normalize_numeric_headings(lines, fence_mask, table_mask)
    normalized = "\n".join(lines)
    if has_trailing_newline:
        normalized += "\n"
    return normalized


def _normalize_table_separators(lines: list[str], fence_mask: list[bool]) -> list[str]:
    updated = list(lines)
    i = 0
    while i + 1 < len(updated):
        if fence_mask[i] or fence_mask[i + 1]:
            i += 1
            continue
        header_line = updated[i]
        sep_line = updated[i + 1]
        if not _is_table_header_line(header_line) or not _is_table_separator_line(sep_line):
            i += 1
            continue
        header_cells = _split_pipe_row(header_line)
        if not header_cells:
            i += 1
            continue
        sep_cells = _split_pipe_row(sep_line)
        sep_cells = _normalize_separator_cells(sep_cells, len(header_cells))
        updated[i + 1] = "|" + " | ".join(sep_cells) + "|"
        i += 2
    return updated


def _normalize_numeric_headings(
    lines: list[str], fence_mask: list[bool], table_mask: list[bool]
) -> list[str]:
    updated: list[str] = []
    in_list = False
    list_indent = 0
    for idx, line in enumerate(lines):
        stripped = line.rstrip("\n")
        if fence_mask[idx] or table_mask[idx] or _is_indented_code_line(stripped):
            updated.append(stripped)
            in_list, list_indent = _update_list_state(stripped, in_list, list_indent)
            continue
        if _should_convert_numeric_heading(lines, idx, in_list):
            match = _NUMERIC_HEADING_RE.match(stripped)
            num = match.group("num") if match else ""
            depth = num.count(".") + 1 if num else 1
            level = min(4, depth + 1)
            updated.append(f"{'#' * level} {stripped}")
            in_list = False
            list_indent = 0
            continue

        updated.append(stripped)
        in_list, list_indent = _update_list_state(stripped, in_list, list_indent)
    return updated


def _should_convert_numeric_heading(lines: list[str], index: int, in_list: bool) -> bool:
    line = lines[index]
    stripped = line.strip("\n")
    if not stripped:
        return False
    if line[:1].isspace():
        return False
    if stripped.startswith("#"):
        return False
    if stripped.lstrip().startswith("|"):
        return False
    if not _NUMERIC_HEADING_RE.match(stripped):
        return False
    if not _is_title_like_numeric_heading(lines, index, stripped):
        return False
    if in_list:
        return False
    if _is_consecutive_ordered_list(lines, index):
        return False
    return True


def _is_title_like_numeric_heading(lines: list[str], index: int, stripped: str) -> bool:
    if len(stripped) > 80:
        return False
    if not _has_required_numeric_marker(stripped):
        return False
    if not _next_line_allows_heading(lines, index):
        return False
    if _REQUIRE_NUMERIC_HEADING_CAPS and not _starts_with_uppercase_alpha(stripped):
        return False
    return True


def _has_required_numeric_marker(text: str) -> bool:
    match = _NUMERIC_HEADING_RE.match(text)
    if not match:
        return False
    num = match.group("num")
    if "." in num:
        return True
    return text.startswith(f"{num}.")

def _next_line_allows_heading(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    next_line = lines[index + 1]
    if not next_line.strip():
        return True
    if next_line.lstrip().startswith("|"):
        return False
    return True


def _starts_with_uppercase_alpha(text: str) -> bool:
    match = _NUMERIC_HEADING_RE.match(text)
    if not match:
        return False
    num = match.group("num")
    marker = f"{num}."
    remainder = text[len(marker) :] if text.startswith(marker) else text[len(num) :]
    for ch in remainder:
        if ch.isalpha():
            return ch.isupper()
    return True


def _is_consecutive_ordered_list(lines: list[str], index: int) -> bool:
    line = lines[index].strip()
    if not _ORDERED_LIST_RE.match(line):
        return False
    if index > 0:
        prev = lines[index - 1].strip()
        if prev and _ORDERED_LIST_RE.match(prev):
            return True
    if index + 1 < len(lines):
        nxt = lines[index + 1].strip()
        if nxt and _ORDERED_LIST_RE.match(nxt):
            return True
    return False


def _update_list_state(line: str, in_list: bool, list_indent: int) -> tuple[bool, int]:
    if not line.strip():
        return False, 0
    if _is_list_marker_line(line):
        return True, _leading_ws(line)
    if in_list and _leading_ws(line) > list_indent:
        return True, list_indent
    return False, 0


def _is_list_marker_line(line: str) -> bool:
    if re.match(r"^\s*[-*+]\s+\S", line):
        return True
    return bool(re.match(r"^\s*\d+[.)]\s+\S", line))


def _leading_ws(line: str) -> int:
    count = 0
    for ch in line:
        if ch == " ":
            count += 1
        elif ch == "\t":
            count += 4
        else:
            break
    return count


def _is_indented_code_line(line: str) -> bool:
    return bool(re.match(r"^(?:\t| {4,})", line))


def _is_table_header_line(line: str) -> bool:
    return line.lstrip().startswith("|")


def _is_table_separator_line(line: str) -> bool:
    stripped = line.strip()
    if "-" not in stripped:
        return False
    return bool(_TABLE_SEP_LINE_RE.match(stripped))


def _split_pipe_row(line: str) -> list[str]:
    text = line.strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|"):
        text = text[:-1]
    return [cell.strip() for cell in text.split("|")]


def _normalize_separator_cells(sep_cells: list[str], target_count: int) -> list[str]:
    normalized: list[str] = []
    for idx in range(target_count):
        raw = sep_cells[idx] if idx < len(sep_cells) else ""
        token = raw.strip().replace(" ", "")
        if token and _TABLE_SEP_CELL_RE.match(token):
            normalized.append(token)
        else:
            normalized.append("---")
    return normalized


def _compute_fence_mask(lines: list[str]) -> list[bool]:
    mask = [False] * len(lines)
    in_fence = False
    for i, line in enumerate(lines):
        if _FENCE_RE.match(line):
            in_fence = not in_fence
            mask[i] = True
            continue
        mask[i] = in_fence
    return mask


def _compute_table_mask(lines: list[str], fence_mask: list[bool]) -> list[bool]:
    mask = [False] * len(lines)
    i = 0
    while i + 1 < len(lines):
        if fence_mask[i] or fence_mask[i + 1]:
            i += 1
            continue
        header_line = lines[i]
        sep_line = lines[i + 1]
        if not _is_table_header_line(header_line) or not _is_table_separator_line(sep_line):
            i += 1
            continue
        header_cells = _split_pipe_row(header_line)
        sep_cells = _split_pipe_row(sep_line)
        if not header_cells or not sep_cells:
            i += 1
            continue
        mask[i] = True
        mask[i + 1] = True
        j = i + 2
        while j < len(lines):
            if fence_mask[j]:
                break
            row_line = lines[j]
            if not row_line.strip() or "|" not in row_line:
                break
            mask[j] = True
            j += 1
        i = j
    return mask


class _DocWriter:
    def __init__(
        self,
        doc: Document,
        insert_after: Paragraph | None,
        reuse_first: Paragraph | None = None,
    ) -> None:
        self.doc = doc
        self.insert_after = insert_after
        self.reuse_first = reuse_first

    def add_paragraph(self, style: str | None = None) -> Paragraph:
        if self.reuse_first is not None:
            paragraph = self.reuse_first
            self.reuse_first = None
        elif self.insert_after is None:
            paragraph = self.doc.add_paragraph()
        else:
            paragraph = _insert_paragraph_after(self.insert_after)
        if style:
            try:
                paragraph.style = style
            except KeyError:
                pass
        self.insert_after = paragraph
        return paragraph

    def add_table(self, rows: int, cols: int):
        table = self.doc.add_table(rows=rows, cols=cols)
        if self.insert_after is not None:
            self.insert_after._p.addnext(table._tbl)
        new_p = OxmlElement("w:p")
        table._tbl.addnext(new_p)
        paragraph_after = Paragraph(new_p, table._parent)
        self.reuse_first = paragraph_after
        self.insert_after = paragraph_after
        return table


class _Token:
    __slots__ = ("type", "tag", "content", "children", "attrs")

    def __init__(
        self,
        token_type: str,
        *,
        tag: str = "",
        content: str = "",
        children: list["_Token"] | None = None,
        attrs: list[tuple[str, str]] | None = None,
    ) -> None:
        self.type = token_type
        self.tag = tag
        self.content = content
        self.children = children or []
        self.attrs = attrs or []


class _FallbackMarkdownParser:
    _table_sep_re = re.compile(r"^:?-{3,}:?$")

    def parse(self, markdown_text: str) -> list[_Token]:
        lines = (markdown_text or "").splitlines()
        tokens: list[_Token] = []
        i = 0
        while i < len(lines):
            line = lines[i]
            if not line.strip():
                i += 1
                continue

            heading = self._parse_heading_line(line)
            if heading is not None:
                level, text = heading
                tokens.extend(self._heading_tokens(level, text))
                i += 1
                continue

            table = self._parse_table(lines, i)
            if table is not None:
                table_tokens, next_i = table
                tokens.extend(table_tokens)
                i = next_i
                continue

            if self._is_bullet_item(line):
                list_tokens, next_i = self._parse_bullet_list(lines, i)
                tokens.extend(list_tokens)
                i = next_i
                continue

            paragraph_text, next_i = self._parse_paragraph(lines, i)
            tokens.extend(self._paragraph_tokens(paragraph_text))
            i = next_i

        return tokens

    def _parse_heading_line(self, line: str) -> tuple[int, str] | None:
        match = re.match(r"^(#{1,3})\s+(.*)$", line.strip())
        if not match:
            return None
        hashes, text = match.groups()
        level = min(3, len(hashes))
        return level, (text or "").strip()

    def _is_bullet_item(self, line: str) -> bool:
        return bool(re.match(r"^\s*[-*]\s+\S", line))

    def _parse_bullet_list(self, lines: list[str], start: int) -> tuple[list[_Token], int]:
        tokens: list[_Token] = [_Token("bullet_list_open")]
        i = start
        while i < len(lines):
            line = lines[i]
            if not line.strip():
                break
            match = re.match(r"^\s*[-*]\s+(.*)$", line)
            if not match:
                break
            item_text = (match.group(1) or "").strip()
            i += 1
            # Continuation lines that are indented and not a new bullet.
            while i < len(lines):
                cont = lines[i]
                if not cont.strip():
                    break
                if self._is_bullet_item(cont) or self._parse_heading_line(cont) is not None:
                    break
                if cont.startswith("  ") or cont.startswith("\t"):
                    item_text = (item_text + " " + cont.strip()).strip()
                    i += 1
                    continue
                break
            tokens.append(_Token("list_item_open"))
            tokens.extend(self._paragraph_tokens(item_text))
            tokens.append(_Token("list_item_close"))
        tokens.append(_Token("bullet_list_close"))
        return tokens, i

    def _parse_paragraph(self, lines: list[str], start: int) -> tuple[str, int]:
        parts: list[str] = []
        i = start
        while i < len(lines):
            line = lines[i]
            if not line.strip():
                break
            if self._parse_heading_line(line) is not None:
                break
            if self._is_bullet_item(line):
                break
            table = self._parse_table(lines, i)
            if table is not None:
                break
            parts.append(line.strip())
            i += 1
        return " ".join(parts).strip(), i

    def _parse_table(self, lines: list[str], start: int) -> tuple[list[_Token], int] | None:
        if start + 1 >= len(lines):
            return None
        header_line = lines[start].rstrip()
        sep_line = lines[start + 1].rstrip()
        if "|" not in header_line or "|" not in sep_line:
            return None
        header_cells = self._split_pipe_row(header_line)
        sep_cells = self._split_pipe_row(sep_line)
        if not header_cells or not sep_cells:
            return None
        if len(header_cells) != len(sep_cells):
            return None
        if not all(self._table_sep_re.match(cell.replace(" ", "")) for cell in sep_cells):
            return None

        rows: list[list[str]] = [header_cells]
        i = start + 2
        while i < len(lines):
            line = lines[i].rstrip()
            if not line.strip():
                break
            if "|" not in line:
                break
            row_cells = self._split_pipe_row(line)
            if not row_cells:
                break
            # Normalize row width to header width.
            if len(row_cells) < len(header_cells):
                row_cells = row_cells + [""] * (len(header_cells) - len(row_cells))
            elif len(row_cells) > len(header_cells):
                row_cells = row_cells[: len(header_cells)]
            rows.append(row_cells)
            i += 1

        tokens = self._table_tokens(rows)
        return tokens, i

    def _split_pipe_row(self, line: str) -> list[str]:
        text = line.strip()
        if text.startswith("|"):
            text = text[1:]
        if text.endswith("|"):
            text = text[:-1]
        return [cell.strip() for cell in text.split("|")]

    def _inline_token(self, text: str) -> _Token:
        children = self._parse_inline_children(text or "")
        if not children:
            children = [_Token("text", content="")]
        return _Token("inline", children=children)

    def _parse_inline_children(self, text: str) -> list[_Token]:
        tokens: list[_Token] = []
        i = 0
        while i < len(text):
            ch = text[i]
            if ch == "`":
                end = text.find("`", i + 1)
                if end == -1:
                    tokens.append(_Token("text", content="`"))
                    i += 1
                    continue
                tokens.append(_Token("code_inline", content=text[i + 1 : end]))
                i = end + 1
                continue

            if ch == "[":
                link = self._parse_link(text, i)
                if link is not None:
                    label, url, next_i = link
                    tokens.append(_Token("link_open", attrs=[("href", url)]))
                    if label:
                        tokens.extend(self._parse_inline_children(label))
                    tokens.append(_Token("link_close"))
                    i = next_i
                    continue

            if text.startswith("**", i) or text.startswith("__", i):
                delim = text[i : i + 2]
                if self._is_emphasis_candidate(text, i, 2):
                    end = self._find_delim(text, i + 2, delim)
                    if end is not None:
                        inner = text[i + 2 : end]
                        if inner:
                            tokens.append(_Token("strong_open"))
                            tokens.extend(self._parse_inline_children(inner))
                            tokens.append(_Token("strong_close"))
                            i = end + 2
                            continue

            if ch in ("*", "_"):
                if self._is_emphasis_candidate(text, i, 1):
                    end = self._find_delim(text, i + 1, ch)
                    if end is not None:
                        inner = text[i + 1 : end]
                        if inner:
                            tokens.append(_Token("em_open"))
                            tokens.extend(self._parse_inline_children(inner))
                            tokens.append(_Token("em_close"))
                            i = end + 1
                            continue

            start = i
            i += 1
            while i < len(text) and text[i] not in "`[*_":
                i += 1
            tokens.append(_Token("text", content=text[start:i]))

        return self._merge_text_tokens(tokens)

    def _parse_link(self, text: str, start: int) -> tuple[str, str, int] | None:
        end_label = text.find("]", start + 1)
        if end_label == -1:
            return None
        if end_label + 1 >= len(text) or text[end_label + 1] != "(":
            return None
        end_url = text.find(")", end_label + 2)
        if end_url == -1:
            return None
        label = text[start + 1 : end_label]
        url = text[end_label + 2 : end_url]
        return label, url, end_url + 1

    def _find_delim(self, text: str, start: int, delim: str) -> int | None:
        end = text.find(delim, start)
        if end == -1:
            return None
        if end == start:
            return None
        return end

    def _is_emphasis_candidate(self, text: str, index: int, delim_len: int) -> bool:
        prev = text[index - 1] if index > 0 else ""
        next_char = text[index + delim_len] if index + delim_len < len(text) else ""
        if prev.isalnum() and next_char.isalnum():
            return False
        return True

    def _merge_text_tokens(self, tokens: list[_Token]) -> list[_Token]:
        merged: list[_Token] = []
        for token in tokens:
            if token.type != "text" or not merged or merged[-1].type != "text":
                merged.append(token)
            else:
                merged[-1].content += token.content
        return merged

    def _heading_tokens(self, level: int, text: str) -> list[_Token]:
        tag = f"h{level}"
        return [
            _Token("heading_open", tag=tag),
            self._inline_token(text),
            _Token("heading_close", tag=tag),
        ]

    def _paragraph_tokens(self, text: str) -> list[_Token]:
        return [
            _Token("paragraph_open", tag="p"),
            self._inline_token(text),
            _Token("paragraph_close", tag="p"),
        ]

    def _table_tokens(self, rows: list[list[str]]) -> list[_Token]:
        tokens: list[_Token] = [_Token("table_open")]
        if not rows:
            tokens.append(_Token("table_close"))
            return tokens
        header = rows[0]
        tokens.append(_Token("tr_open"))
        for cell in header:
            tokens.append(_Token("th_open"))
            tokens.append(self._inline_token(cell))
            tokens.append(_Token("th_close"))
        tokens.append(_Token("tr_close"))
        for row in rows[1:]:
            tokens.append(_Token("tr_open"))
            for cell in row:
                tokens.append(_Token("td_open"))
                tokens.append(self._inline_token(cell))
                tokens.append(_Token("td_close"))
            tokens.append(_Token("tr_close"))
        tokens.append(_Token("table_close"))
        return tokens


def _markdown_parser() -> Any:
    try:
        from markdown_it import MarkdownIt  # type: ignore
    except ImportError:
        return _FallbackMarkdownParser()
    md = MarkdownIt("commonmark", {"breaks": False, "html": False})
    md.enable("table")
    return md


def _is_markdown_it(parser: Any) -> bool:
    return parser.__class__.__name__ == "MarkdownIt"


def _init_document(template_path: Path | None) -> tuple[Document, _DocWriter]:
    if template_path:
        _ensure_template(template_path)
        if template_path.exists() and template_path.is_dir():
            raise ValueError(f"Template path is a directory: {template_path}")
        doc = Document(str(template_path))
        placeholder_paragraph = _find_placeholder_paragraph(doc, PLACEHOLDER)
        if placeholder_paragraph is not None:
            _remove_placeholder_text(placeholder_paragraph, PLACEHOLDER)
            reuse_first = None
            if not placeholder_paragraph.text.strip():
                reuse_first = placeholder_paragraph
            return doc, _DocWriter(
                doc, insert_after=placeholder_paragraph, reuse_first=reuse_first
            )
        return doc, _DocWriter(doc, insert_after=None)
    doc = Document()
    return doc, _DocWriter(doc, insert_after=None)


def _ensure_template(template_path: Path) -> None:
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")



def _find_placeholder_paragraph(doc: Document, placeholder: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            return paragraph
    return None


def _remove_placeholder_text(paragraph: Paragraph, placeholder: str) -> None:
    if placeholder not in paragraph.text:
        return
    paragraph.text = paragraph.text.replace(placeholder, "")


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _save_docx_atomic(doc: Document, path: Path) -> None:
    tmp_path = path.with_suffix(path.suffix + ".tmp")
    doc.save(str(tmp_path))
    tmp_path.replace(path)


def _render_tokens(tokens: list[Any], doc: Document, writer: _DocWriter) -> None:
    list_stack: list[str] = []
    in_list_item = 0
    i = 0
    while i < len(tokens):
        token = tokens[i]
        if token.type == "bullet_list_open":
            list_stack.append("bullet")
            i += 1
            continue
        if token.type == "ordered_list_open":
            list_stack.append("ordered")
            i += 1
            continue
        if token.type in ("bullet_list_close", "ordered_list_close"):
            if list_stack:
                list_stack.pop()
            i += 1
            continue
        if token.type == "list_item_open":
            in_list_item += 1
            i += 1
            continue
        if token.type == "list_item_close":
            in_list_item = max(0, in_list_item - 1)
            i += 1
            continue

        if token.type == "heading_open":
            level = int(token.tag[1]) if token.tag and token.tag.startswith("h") else 1
            inline = tokens[i + 1] if i + 1 < len(tokens) and tokens[i + 1].type == "inline" else None
            runs = _inline_runs(inline)
            style = _heading_style(doc, level)
            _add_paragraph(writer, style, runs)
            i = _skip_to(tokens, i, "heading_close") + 1
            continue

        if token.type == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < len(tokens) and tokens[i + 1].type == "inline" else None
            runs = _inline_runs(inline)
            style = _paragraph_style(doc, list_stack, in_list_item)
            _add_paragraph(
                writer,
                style,
                runs,
                list_indent=in_list_item > 0 and list_stack,
            )
            i = _skip_to(tokens, i, "paragraph_close") + 1
            continue

        if token.type == "hr":
            _add_paragraph(writer, _paragraph_style(doc, [], 0), [])
            i += 1
            continue

        if token.type == "fence":
            _add_code_block(writer, token.content, _paragraph_style(doc, [], 0))
            i += 1
            continue

        if token.type == "table_open":
            table_rows, end_index = _parse_table(tokens, i)
            if table_rows:
                table = _render_table(writer, doc, table_rows)
                _apply_table_profile(table, doc)
            i = end_index + 1
            continue

        i += 1


def _skip_to(tokens: list[Any], start_index: int, end_type: str) -> int:
    i = start_index + 1
    while i < len(tokens) and tokens[i].type != end_type:
        i += 1
    return i


def _heading_style(doc: Document, level: int) -> str | None:
    if level <= 1:
        return "Title" if _style_exists(doc, "Title") else "Heading 1"
    if level == 2:
        return "Heading 1"
    return "Heading 2"


def _paragraph_style(doc: Document, list_stack: list[str], in_list_item: int) -> str | None:
    if in_list_item > 0 and list_stack:
        list_kind = list_stack[-1]
        if list_kind == "bullet":
            style = "List Bullet"
        else:
            style = "List Number"
        return style
    return "Normal"


def _style_exists(doc: Document, name: str) -> bool:
    try:
        doc.styles[name]
    except KeyError:
        return False
    return True


def _add_paragraph(
    writer: _DocWriter,
    style: str | None,
    runs: list[dict[str, Any]],
    *,
    list_indent: bool = False,
) -> None:
    paragraph = writer.add_paragraph(style)
    if list_indent:
        _apply_list_indent(paragraph)
    _add_runs(paragraph, runs)


def _add_runs(paragraph: Paragraph, runs: list[dict[str, Any]]) -> None:
    i = 0
    while i < len(runs):
        run_spec = runs[i]
        if run_spec.get("break"):
            paragraph.add_run().add_break()
            i += 1
            continue
        link_url = run_spec.get("link_url")
        if link_url:
            group: list[dict[str, Any]] = []
            while i < len(runs):
                current = runs[i]
                if current.get("break") or current.get("link_url") != link_url:
                    break
                group.append(current)
                i += 1
            if _add_hyperlink_runs(paragraph, group, link_url):
                continue
            for item in group:
                _add_plain_run(paragraph, item)
            paragraph.add_run(f" ({link_url})")
            continue

        _add_plain_run(paragraph, run_spec)
        i += 1


def _add_plain_run(paragraph: Paragraph, run_spec: dict[str, Any]) -> None:
    text = run_spec.get("text", "")
    if not text:
        return
    run = paragraph.add_run(text)
    if run_spec.get("bold"):
        run.bold = True
    if run_spec.get("italic"):
        run.italic = True
    if run_spec.get("code"):
        run.font.name = "Consolas"


def _add_hyperlink_runs(
    paragraph: Paragraph, run_specs: list[dict[str, Any]], link_url: str
) -> bool:
    if not link_url:
        return False
    try:
        part = paragraph.part
        rel_id = part.relate_to(link_url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    except Exception:
        return False

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    for run_spec in run_specs:
        text = run_spec.get("text", "")
        if not text:
            continue
        run = OxmlElement("w:r")
        r_pr = _build_run_properties(run_spec)
        if r_pr is not None:
            run.append(r_pr)
        text_el = OxmlElement("w:t")
        text_el.text = text
        run.append(text_el)
        hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return True


def _build_run_properties(run_spec: dict[str, Any]) -> OxmlElement | None:
    r_pr = OxmlElement("w:rPr")
    if run_spec.get("bold"):
        r_pr.append(OxmlElement("w:b"))
    if run_spec.get("italic"):
        r_pr.append(OxmlElement("w:i"))
    if run_spec.get("code"):
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Consolas")
        r_fonts.set(qn("w:hAnsi"), "Consolas")
        r_pr.append(r_fonts)
    if len(r_pr) == 0:
        return None
    return r_pr


def _add_code_block(writer: _DocWriter, content: str, style: str | None) -> None:
    lines = content.rstrip("\n").splitlines()
    if not lines:
        lines = [""]
    for line in lines:
        paragraph = writer.add_paragraph(style)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"


def _inline_runs(inline_token: Any | None) -> list[dict[str, Any]]:
    if inline_token is None or not getattr(inline_token, "children", None):
        return []
    runs: list[dict[str, Any]] = []
    bold = False
    italic = False
    link_url: str | None = None
    for child in inline_token.children:
        token_type = child.type
        if token_type == "strong_open":
            bold = True
            continue
        if token_type == "strong_close":
            bold = False
            continue
        if token_type == "em_open":
            italic = True
            continue
        if token_type == "em_close":
            italic = False
            continue
        if token_type == "link_open":
            link_url = _extract_link_url(child)
            continue
        if token_type == "link_close":
            link_url = None
            continue
        if token_type == "code_inline":
            _append_run(runs, child.content, False, False, True, link_url)
            continue
        if token_type == "text":
            _append_run(runs, child.content, bold, italic, False, link_url)
            continue
        if token_type == "softbreak":
            _append_run(runs, " ", bold, italic, False, link_url)
            continue
        if token_type == "hardbreak":
            runs.append({"break": True})
            continue
        if token_type == "image":
            alt_text = child.content or ""
            placeholder = "TODO: image omitted"
            if alt_text:
                placeholder = f"TODO: image omitted ({alt_text})"
            _append_run(runs, placeholder, bold, italic, False, link_url)
            continue
        if token_type in ("link_open", "link_close"):
            continue
    return runs


def _append_run(
    runs: list[dict[str, Any]],
    text: str,
    bold: bool,
    italic: bool,
    code: bool,
    link_url: str | None,
) -> None:
    if not text:
        return
    if runs:
        last = runs[-1]
        if (
            not last.get("break")
            and last.get("bold") == bold
            and last.get("italic") == italic
            and last.get("code") == code
            and last.get("link_url") == link_url
        ):
            last["text"] = f"{last.get('text', '')}{text}"
            return
    runs.append(
        {
            "text": text,
            "bold": bold,
            "italic": italic,
            "code": code,
            "link_url": link_url,
        }
    )


def _extract_link_url(token: Any) -> str | None:
    attrs = getattr(token, "attrs", None)
    if not attrs:
        return None
    try:
        return dict(attrs).get("href")
    except Exception:
        return None


def _apply_placeholder_replacements(doc: Document, md_file: Path) -> None:
    company_name = _resolve_company_name(md_file)
    replacements = {
        "{{COMPANY}}": company_name or "Unknown Company",
        "{{DATE}}": date.today().isoformat(),
    }

    for paragraph in list(doc.paragraphs):
        if "{{CONTACT}}" in paragraph.text:
            _remove_paragraph(paragraph)
            continue
        _replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in list(cell.paragraphs):
                    if "{{CONTACT}}" in paragraph.text:
                        _remove_paragraph(paragraph)
                        continue
                    _replace_text_in_paragraph(paragraph, replacements)


def _replace_text_in_paragraph(paragraph: Paragraph, replacements: dict[str, str]) -> None:
    text = paragraph.text
    if not text:
        return
    updated = text
    for key, value in replacements.items():
        if key in updated:
            updated = updated.replace(key, value)
    if updated != text:
        paragraph.text = updated


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._p
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _resolve_company_name(md_file: Path) -> str | None:
    candidate_paths: list[Path] = []
    for parent in [md_file.parent, *md_file.parents]:
        candidate_paths.append(parent / "lead_input.json")
        candidate_paths.append(parent / "_dossier" / "lead_input.json")

    seen: set[Path] = set()
    for path in candidate_paths:
        if path in seen:
            continue
        seen.add(path)
        if not path.exists():
            continue
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        if not isinstance(payload, dict):
            continue
        for key in ("company_name", "client", "company"):
            value = payload.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()
    return None


def _apply_style_profile(doc: Document) -> None:
    profiles = {
        "Title": {
            "font_size_pt": 18.0,
            "bold": True,
            "spacing_before_pt": 24.0,
            "spacing_after_pt": 12.0,
            "keep_with_next": True,
        },
        "Heading 1": {
            "font_size_pt": 16.0,
            "bold": True,
            "spacing_before_pt": 24.0,
            "spacing_after_pt": 0.0,
            "keep_with_next": True,
        },
        "Heading 2": {
            "font_size_pt": 14.0,
            "bold": True,
            "spacing_before_pt": 10.0,
            "spacing_after_pt": 0.0,
            "keep_with_next": True,
        },
        "Heading 3": {
            "bold": True,
            "spacing_before_pt": 10.0,
            "spacing_after_pt": 0.0,
            "keep_with_next": True,
        },
        "Normal": {
            "spacing_before_pt": 9.0,
            "spacing_after_pt": 9.0,
        },
    }

    for style_name, profile in profiles.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        font = style.font
        if profile.get("font_size_pt") is not None and font.size is None:
            font.size = Pt(profile["font_size_pt"])
        if profile.get("bold") is not None and font.bold is None:
            font.bold = profile["bold"]
        fmt = style.paragraph_format
        if profile.get("spacing_before_pt") is not None and fmt.space_before is None:
            fmt.space_before = Pt(profile["spacing_before_pt"])
        if profile.get("spacing_after_pt") is not None and fmt.space_after is None:
            fmt.space_after = Pt(profile["spacing_after_pt"])
        if profile.get("keep_with_next") is not None and fmt.keep_with_next is None:
            fmt.keep_with_next = profile["keep_with_next"]


def _apply_list_indent(paragraph: Paragraph) -> None:
    fmt = paragraph.paragraph_format
    if fmt.left_indent is None:
        fmt.left_indent = Pt(18.0)
    if fmt.first_line_indent is None:
        fmt.first_line_indent = Pt(-9.0)


def _apply_table_profile(table, doc: Document) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    section = doc.sections[0]
    total_width_emu = int(section.page_width - section.left_margin - section.right_margin)
    total_twips = max(1, int(round(total_width_emu / 635.0)))

    col_count = len(table.columns) if table.columns else 0
    if col_count <= 0:
        return

    widths_twips = _compute_table_col_widths_twips(table, total_twips, col_count)
    widths_emu = [max(1, int(w * 635)) for w in widths_twips]

    _set_tbl_grid(table, widths_twips)
    for idx, col in enumerate(table.columns):
        col.width = widths_emu[idx]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            try:
                cell.width = widths_emu[idx]
            except Exception:
                pass
            _tighten_table_cell(cell)

    if table.rows:
        _set_header_repeat(table.rows[0])


def _compute_table_col_widths_twips(table, total_twips: int, col_count: int) -> list[int]:
    if col_count == 1:
        return [total_twips]
    if col_count == 2 and table.rows:
        left = (table.cell(0, 0).text or "").strip().lower()
        right = (table.cell(0, 1).text or "").strip().lower()
        if left == "name" and right == "description":
            first = int(total_twips * 0.30)
            return [first, total_twips - first]
        first = total_twips // 2
        return [first, total_twips - first]
    base = total_twips // col_count
    widths = [base] * col_count
    widths[-1] += total_twips - (base * col_count)
    return widths


def _set_tbl_grid(table, widths_twips: list[int]) -> None:
    tbl = table._tbl
    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        # Ensure tblGrid is placed after tblPr.
        insert_at = 1 if tbl.tblPr is not None else 0
        tbl.insert(insert_at, tbl_grid)
    else:
        for child in list(tbl_grid):
            tbl_grid.remove(child)
    for width in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(max(1, int(width))))
        tbl_grid.append(grid_col)


def _tighten_table_cell(cell) -> None:
    # Remove trailing empty paragraphs and force compact spacing.
    while len(cell.paragraphs) > 1 and not (cell.paragraphs[-1].text or "").strip():
        _remove_paragraph(cell.paragraphs[-1])
    for paragraph in cell.paragraphs:
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)


def _set_header_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def _parse_table(tokens: list[Any], start_index: int) -> tuple[list[list[list[dict[str, Any]]]], int]:
    rows: list[list[list[dict[str, Any]]]] = []
    row_cells: list[list[dict[str, Any]]] = []
    i = start_index + 1
    while i < len(tokens):
        token = tokens[i]
        if token.type == "tr_open":
            row_cells = []
        elif token.type in ("th_open", "td_open"):
            runs: list[dict[str, Any]] = []
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                runs = _inline_runs(tokens[i + 1])
            row_cells.append(runs)
        elif token.type == "tr_close":
            rows.append(row_cells)
        elif token.type == "table_close":
            return rows, i
        i += 1
    return rows, i


def _render_table(
    writer: _DocWriter, doc: Document, rows: list[list[list[dict[str, Any]]]]
):
    if not rows:
        return None
    cols = max(len(row) for row in rows)
    table = writer.add_table(rows=len(rows), cols=cols)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            runs = row[c_idx] if c_idx < len(row) else []
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            _add_runs(paragraph, runs)
    return table
