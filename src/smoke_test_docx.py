from __future__ import annotations

import argparse
import json
import os
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from dotenv import find_dotenv, load_dotenv

from primer_ops.render_docx import render_primer_docx
from primer_ops.primer import resolve_lead_input_path, resolve_output_targets
from scripts.compare_docx_structure import compare_docx_structure


def _style_exists(doc: Document, name: str) -> bool:
    try:
        doc.styles[name]
    except KeyError:
        return False
    return True


def _get_paragraph_by_text(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise AssertionError(f"Paragraph not found: {text}")


def _assert_heading_styles() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        md_path = tmp_path / "sample.md"
        docx_path = tmp_path / "sample.docx"
        md_path.write_text(
            "# Title\n\n## Section\n\n### Subsection\n\nParagraph.\n", encoding="utf-8"
        )
        render_primer_docx(str(md_path), str(docx_path), None)
        doc = Document(str(docx_path))

        title_para = _get_paragraph_by_text(doc, "Title")
        section_para = _get_paragraph_by_text(doc, "Section")
        subsection_para = _get_paragraph_by_text(doc, "Subsection")

        expected_title = "Title" if _style_exists(doc, "Title") else "Heading 1"
        if _style_exists(doc, expected_title):
            assert (
                title_para.style.name == expected_title
            ), f"Title style mismatch: {title_para.style.name}"

        if _style_exists(doc, "Heading 1"):
            assert (
                section_para.style.name == "Heading 1"
            ), f"Heading 1 style mismatch: {section_para.style.name}"

        if _style_exists(doc, "Heading 2"):
            assert (
                subsection_para.style.name == "Heading 2"
            ), f"Heading 2 style mismatch: {subsection_para.style.name}"


def _assert_inline_markdown_render() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        md_path = tmp_path / "inline.md"
        docx_path = tmp_path / "inline.docx"
        md_path.write_text(
            "Paragraph with **bold** and *italic* and [site](https://example.com) "
            "plus ([site](https://example.com)).\n",
            encoding="utf-8",
        )
        render_primer_docx(str(md_path), str(docx_path), None)
        doc = Document(str(docx_path))
        full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        assert "**" not in full_text, "Raw bold markers leaked into DOCX text."
        assert "](" not in full_text, "Raw link markers leaked into DOCX text."
        doc_xml = doc.element.xml
        assert (
            "site" in full_text or "site" in doc_xml
        ), "Expected link label text missing."

        has_url_text = "https://example.com" in full_text
        has_hyperlink_rel = any(
            rel.reltype == RELATIONSHIP_TYPE.HYPERLINK
            and rel.target_ref == "https://example.com"
            for rel in doc.part.rels.values()
        )
        assert has_url_text or has_hyperlink_rel, "Expected link URL not found."


def _assert_numeric_heading_normalization() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        md_path = tmp_path / "numeric.md"
        docx_path = tmp_path / "numeric.docx"
        md_path.write_text(
            "1. Company Introduction\n\n2.1 Ownership & Governance\n\nParagraph.\n",
            encoding="utf-8",
        )
        render_primer_docx(str(md_path), str(docx_path), None)
        doc = Document(str(docx_path))
        intro_para = _get_paragraph_by_text(doc, "1. Company Introduction")
        ownership_para = _get_paragraph_by_text(doc, "2.1 Ownership & Governance")

        if _style_exists(doc, "Heading 1"):
            assert (
                intro_para.style.name == "Heading 1"
            ), f"Numeric heading styled as {intro_para.style.name}"
        else:
            assert intro_para.style.name != "List Number"

        if _style_exists(doc, "Heading 2"):
            assert (
                ownership_para.style.name == "Heading 2"
            ), f"Numeric heading styled as {ownership_para.style.name}"
        else:
            assert ownership_para.style.name != "List Number"


def _resolve_latest_md_path(lead_input: str | None, output_dir: str | None) -> Path:
    lead_path = resolve_lead_input_path(lead_input)
    if not lead_path.exists():
        raise SystemExit(
            f"ERROR: lead_input.json not found at {lead_path}. "
            "Use --lead-input or set LEAD_INPUT_PATH."
        )
    lead = json.loads(lead_path.read_text(encoding="utf-8"))
    if not isinstance(lead, dict):
        raise SystemExit("ERROR: lead_input.json must contain a JSON object.")
    targets = resolve_output_targets(output_dir, lead)
    return targets["output_dir"] / "primer.md"


def main() -> None:
    parser = argparse.ArgumentParser(description="DOCX smoke test")
    parser.add_argument(
        "--lead-input",
        default=None,
        help="Path to lead_input.json (default: ./lead_input.json or LEAD_INPUT_PATH)",
    )
    parser.add_argument(
        "--output-dir",
        default=None,
        help="Override OUTPUT_BASE_DIR/OUTPUT_DIR (use as final output folder)",
    )
    args = parser.parse_args()

    load_dotenv(find_dotenv(usecwd=True), override=False)
    _assert_heading_styles()
    _assert_inline_markdown_render()
    _assert_numeric_heading_normalization()
    md_path = _resolve_latest_md_path(args.lead_input, args.output_dir)
    docx_path = md_path.with_suffix(".docx")
    template_path = os.getenv("PRIMER_WORD_TEMPLATE_PATH", "").strip() or None

    if not md_path.exists():
        print(f"DOCX render skipped/failed: {md_path} not found")
        return

    try:
        render_primer_docx(str(md_path), str(docx_path), template_path)
    except Exception as err:
        print(f"DOCX render skipped/failed: {err}")
        return

    print(f"DOCX generated: {docx_path}")

    golden_path = Path("docs") / "golden" / "primer_reference.docx"
    if not golden_path.exists():
        print(f"DOCX compare skipped/failed: {golden_path} not found")
        return

    report_path = Path("out") / "docx_compare_report.md"
    result = compare_docx_structure(docx_path, golden_path, report_path)
    print(f"DOCX compare report: {result.report_path}")
    if result.major_mismatches:
        print("Major mismatches detected:")
        for item in result.major_mismatches:
            print(f"- {item}")
        sys.exit(3)


if __name__ == "__main__":
    main()
